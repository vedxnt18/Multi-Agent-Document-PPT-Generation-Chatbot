"""
backend/services/document_service.py

Extraction for PDF, DOCX, and standalone image files, producing the
normalized ExtractionResult schema.

PDF strategy (per assignment requirement — don't assume every PDF has
selectable text):
    1. Try native text extraction per page (PyMuPDF).
    2. If a page's native text is below a minimal length threshold, treat
       it as likely scanned/image-only and fall back to OCR on a rendered
       image of that page.
    3. Track which pages used OCR vs native text, and surface warnings.

DOCX strategy:
    Walk paragraphs in document order, classify headings vs body via the
    paragraph style name, and extract tables separately.
"""
import logging
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from PIL import Image

from backend.schemas.extraction import ExtractionResult, ContentBlock, BlockType, SourceType
from backend.services.ocr_service import ocr_service

logger = logging.getLogger(__name__)

# If native-extracted text on a PDF page is shorter than this, we assume the
# page is scanned/image-only and attempt OCR instead.
MIN_NATIVE_TEXT_CHARS = 20

# Render resolution for OCR fallback (higher = better OCR accuracy, slower)
OCR_RENDER_ZOOM = 2.0


def extract_pdf(file_path: str, file_id: str, original_filename: str) -> ExtractionResult:
    result = ExtractionResult(
        file_id=file_id,
        original_filename=original_filename,
        file_extension=".pdf",
        document_type="pdf",
    )

    try:
        doc = fitz.open(file_path)
    except Exception as e:
        result.warnings.append(f"Failed to open PDF: {e}")
        return result

    result.page_or_slide_count = doc.page_count

    for page_index in range(doc.page_count):
        page = doc[page_index]
        page_number = page_index + 1

        native_text = page.get_text("text").strip()

        if len(native_text) >= MIN_NATIVE_TEXT_CHARS:
            result.blocks.append(
                ContentBlock(
                    block_type=BlockType.PARAGRAPH,
                    text=native_text,
                    page_number=page_number,
                    source_type=SourceType.NATIVE_TEXT,
                )
            )
        else:
            # Likely scanned page — attempt OCR on a rendered image of it.
            if not ocr_service.available:
                result.warnings.append(
                    f"Page {page_number}: insufficient selectable text and OCR is unavailable. "
                    "Unable to extract text from this page."
                )
                continue

            pix = page.get_pixmap(matrix=fitz.Matrix(OCR_RENDER_ZOOM, OCR_RENDER_ZOOM))
            image = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
            ocr_result = ocr_service.extract_text_from_image(image)

            if ocr_result.success:
                result.used_ocr = True
                result.blocks.append(
                    ContentBlock(
                        block_type=BlockType.PARAGRAPH,
                        text=ocr_result.text,
                        page_number=page_number,
                        source_type=SourceType.OCR,
                        confidence=ocr_result.confidence,
                    )
                )
                result.warnings.append(f"Page {page_number}: extracted via OCR (no selectable text found).")
            else:
                result.warnings.append(
                    f"Page {page_number}: unable to extract text (OCR failed: {ocr_result.error})"
                )

    doc.close()
    return result


# Heuristic mapping from python-docx paragraph style names to heading levels.
def _heading_level_from_style(style_name: str) -> int | None:
    if not style_name:
        return None
    name = style_name.lower()
    if name.startswith("heading"):
        parts = name.split(" ")
        if len(parts) == 2 and parts[1].isdigit():
            return int(parts[1])
        return 1
    if name in ("title",):
        return 0
    return None


def extract_docx(file_path: str, file_id: str, original_filename: str) -> ExtractionResult:
    result = ExtractionResult(
        file_id=file_id,
        original_filename=original_filename,
        file_extension=".docx",
        document_type="docx",
    )

    try:
        doc = DocxDocument(file_path)
    except Exception as e:
        result.warnings.append(f"Failed to open DOCX: {e}")
        return result

    # Paragraphs (headings + body text), in document order
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        level = _heading_level_from_style(para.style.name if para.style else "")
        block_type = BlockType.HEADING if level is not None else BlockType.PARAGRAPH

        # Capture basic run-level formatting from the first run as a signal
        # for style analysis in Phase 3 (not exhaustive — just what's cheap here).
        font_name = None
        font_size = None
        alignment = str(para.alignment) if para.alignment is not None else None
        if para.runs:
            first_run = para.runs[0]
            font_name = first_run.font.name
            font_size = first_run.font.size.pt if first_run.font.size else None

        result.blocks.append(
            ContentBlock(
                block_type=block_type,
                text=text,
                level=level,
                source_type=SourceType.NATIVE_TEXT,
                metadata={
                    "style_name": para.style.name if para.style else None,
                    "font_name": font_name,
                    "font_size_pt": font_size,
                    "alignment": alignment,
                },
            )
        )

    # Tables
    for t_idx, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        result.blocks.append(
            ContentBlock(
                block_type=BlockType.TABLE,
                text=f"[Table {t_idx + 1} with {len(rows)} rows]",
                source_type=SourceType.NATIVE_TEXT,
                metadata={"rows": rows},
            )
        )

    result.page_or_slide_count = 0  # DOCX has no reliable page count without rendering
    return result


def extract_image(file_path: str, file_id: str, original_filename: str) -> ExtractionResult:
    ext = Path(file_path).suffix.lower()
    result = ExtractionResult(
        file_id=file_id,
        original_filename=original_filename,
        file_extension=ext,
        document_type="image",
        page_or_slide_count=1,
    )

    if not ocr_service.available:
        result.warnings.append(
            "OCR is unavailable (Tesseract not found). Unable to extract text from this image."
        )
        return result

    try:
        image = Image.open(file_path).convert("RGB")
    except Exception as e:
        result.warnings.append(f"Failed to open image: {e}")
        return result

    ocr_result = ocr_service.extract_text_from_image(image)
    result.used_ocr = True

    if ocr_result.success:
        result.blocks.append(
            ContentBlock(
                block_type=BlockType.PARAGRAPH,
                text=ocr_result.text,
                page_number=1,
                source_type=SourceType.OCR,
                confidence=ocr_result.confidence,
            )
        )
    else:
        result.warnings.append(f"Unable to extract text from this image (OCR failed: {ocr_result.error})")

    return result
