"""
backend/agents/document_generator.py

Document Generation Agent. Consumes a ContentPlan (Phase 8's content
planner) and an optional DocumentStyleSpec (from Phase 3's Document
Analysis Agent, if a template was uploaded) and produces a real, editable
.docx file via python-docx.

Style application (what we can and can't reliably transfer, stated
honestly rather than pretending pixel-perfect cloning):
    - Primary font: applied to the Normal style if detected.
    - Heading levels: mapped to python-docx's built-in Heading 1/2/3 styles.
    - Section order: sections are written in the order the content plan
      gives them (which itself can be informed by the template's
      section_titles order — that's the caller's responsibility).
    - Tables: written as real python-docx tables, not images.
    - Citations: appended inline as bracketed IDs (e.g. "... [WEB-001]"),
      plus a "Sources" appendix section listing every citation used.

We do NOT attempt to clone exact margins/colors/spacing byte-for-byte from
the source document — python-docx doesn't expose a reliable "copy style"
primitive across arbitrary templates, so instead we apply the *facts* the
style spec gives us (primary font, heading structure) which is the
practical, honest interpretation of "preserve style."
"""
import logging
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from backend.schemas.content_plan import ContentPlan
from backend.schemas.template_spec import DocumentStyleSpec
from backend.services.citation_service import citation_registry
from backend.config import settings

logger = logging.getLogger(__name__)


def _apply_primary_font(doc: Document, font_name: str | None) -> None:
    if not font_name:
        return
    try:
        normal_style = doc.styles["Normal"]
        normal_style.font.name = font_name
    except Exception as e:
        logger.warning(f"Could not apply primary font '{font_name}': {e}")


def _add_sources_appendix(doc: Document, citation_ids_used: set[str]) -> None:
    if not citation_ids_used:
        return
    doc.add_page_break()
    doc.add_heading("Sources", level=1)
    for cid in sorted(citation_ids_used):
        source = citation_registry.get(cid)
        p = doc.add_paragraph()
        if source:
            run = p.add_run(f"[{cid}] {source.title}")
            run.bold = True
            if source.detail:
                p.add_run(f" — {source.detail}")
        else:
            p.add_run(f"[{cid}] (source details unavailable)")


def generate_docx(
    content_plan: ContentPlan,
    output_path: str,
    style_spec: DocumentStyleSpec | None = None,
) -> str:
    """
    Writes a real .docx file to output_path. Returns the path written.
    """
    doc = Document()

    if style_spec and style_spec.primary_font:
        _apply_primary_font(doc, style_spec.primary_font)

    # --- Title ---
    title_heading = doc.add_heading(content_plan.title, level=0)
    if content_plan.subtitle:
        subtitle_p = doc.add_paragraph(content_plan.subtitle)
        subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in subtitle_p.runs:
            run.italic = True

    citation_ids_used: set[str] = set()

    # --- Sections ---
    for section in content_plan.sections:
        level = max(1, min(section.level, 4))
        doc.add_heading(section.heading, level=level)

        for para_text in section.paragraphs:
            doc.add_paragraph(para_text)

        for bullet in section.bullet_points:
            doc.add_paragraph(bullet, style="List Bullet")

        if section.table:
            table = doc.add_table(rows=1, cols=len(section.table.headers))
            table.style = "Light Grid Accent 1"
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(section.table.headers):
                hdr_cells[i].text = header
            for row_data in section.table.rows:
                row_cells = table.add_row().cells
                for i, cell_value in enumerate(row_data):
                    if i < len(row_cells):
                        row_cells[i].text = str(cell_value)

        if section.citation_ids:
            citation_line = doc.add_paragraph()
            citation_run = citation_line.add_run("Sources: " + ", ".join(f"[{c}]" for c in section.citation_ids))
            citation_run.italic = True
            citation_run.font.size = Pt(9)
            citation_ids_used.update(section.citation_ids)

    _add_sources_appendix(doc, citation_ids_used)

    output_dir = Path(output_path).parent
    output_dir.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)

    return output_path


def default_output_path(artifact_id: str, version: int = 1) -> str:
    generated_dir = settings.resolve_path(settings.generated_dir)
    return str(generated_dir / f"{artifact_id}_v{version}.docx")
