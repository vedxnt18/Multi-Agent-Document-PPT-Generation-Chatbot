"""
tests/test_conversion.py

Phase 12 tests: verifies conversion is a structural transform (paragraphs
<-> bullets), not raw text copying, and that both directions produce real,
openable, validated files as new artifacts (original preserved).
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

import pytest
from docx import Document as DocxDocument
from pptx import Presentation
from fastapi.testclient import TestClient

from backend.main import app
from backend.agents.converter import docx_plan_to_pptx_plan, pptx_plan_to_docx_plan
from backend.schemas.content_plan import ContentPlan, GeneratedSection
from backend.services.citation_service import citation_registry

client = TestClient(app)


@pytest.fixture(autouse=True)
def _reset():
    citation_registry.reset()
    yield


def test_docx_plan_to_pptx_plan_converts_paragraphs_to_bullets():
    plan = ContentPlan(title="T", sections=[GeneratedSection(heading="Market Overview", paragraphs=["The market is growing fast. It saw huge gains."])])
    converted = docx_plan_to_pptx_plan(plan)
    assert converted.sections[0].bullet_points
    assert not converted.sections[0].paragraphs


def test_pptx_plan_to_docx_plan_converts_bullets_to_prose():
    plan = ContentPlan(title="T", sections=[GeneratedSection(heading="Key Points", bullet_points=["Point one", "Point two"])])
    converted = pptx_plan_to_docx_plan(plan)
    assert converted.sections[0].paragraphs
    assert not converted.sections[0].bullet_points
    assert "Point one" in converted.sections[0].paragraphs[0]
    assert "Point two" in converted.sections[0].paragraphs[0]


def test_convert_document_to_presentation_api():
    gen_resp = client.post("/generate/document", json={"request": "Create a proposal on AI"})
    artifact_id = gen_resp.json()["artifact_id"]

    conv_resp = client.post("/convert/document-to-presentation", json={"artifact_id": artifact_id, "slide_count": 4})
    assert conv_resp.status_code == 200
    body = conv_resp.json()
    assert body["new_artifact_id"] != artifact_id

    prs = Presentation(body["file_path"])
    assert len(prs.slides) == 4

    # original artifact untouched
    orig = client.get(f"/artifact/{artifact_id}")
    assert orig.json()["current_version"] == 1


def test_convert_presentation_to_document_api():
    gen_resp = client.post("/generate/presentation", json={"request": "Create a deck", "slide_count": 3})
    artifact_id = gen_resp.json()["artifact_id"]

    conv_resp = client.post("/convert/presentation-to-document", json={"artifact_id": artifact_id})
    assert conv_resp.status_code == 200
    body = conv_resp.json()

    doc = DocxDocument(body["file_path"])
    assert len(doc.paragraphs) > 0


def test_convert_rejects_wrong_source_type():
    gen_resp = client.post("/generate/document", json={"request": "Create a report"})
    artifact_id = gen_resp.json()["artifact_id"]

    resp = client.post("/convert/presentation-to-document", json={"artifact_id": artifact_id})
    assert resp.status_code == 400


def test_convert_nonexistent_artifact_404():
    resp = client.post("/convert/document-to-presentation", json={"artifact_id": "nope"})
    assert resp.status_code == 404
