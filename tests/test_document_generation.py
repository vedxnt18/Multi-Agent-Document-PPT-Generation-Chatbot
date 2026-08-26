"""
tests/test_document_generation.py

Phase 8 tests. Verifies:
  - content planner produces a valid ContentPlan (heuristic path, since
    LLM_PROVIDER=mock in this environment)
  - generate_docx produces a REAL, openable .docx file with the expected
    structure (headings, paragraphs, tables, citations) — verified by
    re-opening the generated file with python-docx, not just checking
    that a file exists.
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

import pytest
from docx import Document as DocxDocument
from fastapi.testclient import TestClient

from backend.main import app
from backend.agents.content_planner import create_content_plan
from backend.agents.document_generator import generate_docx
from backend.schemas.content_plan import ContentPlan, GeneratedSection, GeneratedTable
from backend.schemas.template_spec import DocumentStyleSpec
from backend.services.citation_service import citation_registry

client = TestClient(app)


@pytest.fixture(autouse=True)
def _reset_citations():
    citation_registry.reset()
    yield


def test_content_planner_heuristic_with_research_findings():
    findings = ["AI adoption grew 40% in 2026 [WEB-001].", "Enterprise spending on GenAI doubled [WEB-002]."]
    plan = create_content_plan("Create a report on AI trends", research_findings=findings)
    assert plan.sections
    assert any("Research Findings" in s.heading for s in plan.sections)
    assert "WEB-001" in plan.sections[0].citation_ids


def test_content_planner_heuristic_with_rag_chunks():
    citation_registry.reset()
    cid = citation_registry.register("RAG", "internal_doc.docx", "chunk 0")
    chunks = [{"citation_id": cid, "text": "Internal revenue grew 15% last quarter."}]
    plan = create_content_plan("Summarize our internal performance", rag_chunks=chunks)
    assert any("Enterprise Knowledge" in s.heading for s in plan.sections)
    assert cid in plan.sections[-1].citation_ids


def test_generate_docx_produces_real_openable_file(tmp_path):
    citation_registry.reset()
    cid = citation_registry.register("WEB", "Example Source", "https://example.com/article")

    plan = ContentPlan(
        title="Test Proposal",
        subtitle="A test subtitle",
        sections=[
            GeneratedSection(
                heading="Executive Summary",
                level=1,
                paragraphs=["This is the executive summary."],
                citation_ids=[cid],
            ),
            GeneratedSection(
                heading="Market Data",
                level=1,
                bullet_points=["Point one", "Point two"],
                table=GeneratedTable(headers=["Quarter", "Revenue"], rows=[["Q1", "$1M"], ["Q2", "$1.2M"]]),
            ),
        ],
    )

    output_path = str(tmp_path / "test_output.docx")
    result_path = generate_docx(plan, output_path)

    assert Path(result_path).exists()

    # Re-open with python-docx to verify real structure, not just file existence
    doc = DocxDocument(result_path)
    full_text = "\n".join(p.text for p in doc.paragraphs)

    assert "Test Proposal" in full_text
    assert "Executive Summary" in full_text
    assert "Market Data" in full_text
    assert "Point one" in full_text
    assert cid in full_text  # inline citation
    assert "Sources" in full_text  # appendix

    assert len(doc.tables) == 1
    assert doc.tables[0].cell(0, 0).text == "Quarter"
    assert doc.tables[0].cell(1, 1).text == "$1M"


def test_generate_docx_applies_primary_font(tmp_path):
    plan = ContentPlan(title="Font Test", sections=[GeneratedSection(heading="Section", paragraphs=["text"])])
    style_spec = DocumentStyleSpec(file_id="x", primary_font="Georgia")

    output_path = str(tmp_path / "font_test.docx")
    generate_docx(plan, output_path, style_spec=style_spec)

    doc = DocxDocument(output_path)
    assert doc.styles["Normal"].font.name == "Georgia"


def test_generate_document_api_end_to_end():
    resp = client.post(
        "/generate/document",
        json={
            "request": "Create a short report on AI adoption",
            "research_findings": ["AI adoption is rising fast [WEB-001]."],
        },
    )
    assert resp.status_code == 200
    body = resp.json()
    assert body["filename"].endswith(".docx")
    assert Path(body["file_path"]).exists()

    # Re-open the actual generated file via the download endpoint
    download_resp = client.get(f"/generate/download/{body['artifact_id']}?ext=docx")
    assert download_resp.status_code == 200
